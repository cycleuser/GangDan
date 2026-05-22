"""Generate Word documents for GangDan-Refined (English and Chinese).

Creates properly formatted .docx files with:
- Professional heading hierarchy
- Architecture diagrams as tables/figures
- Agent reference tables
- Code examples in monospace
- Proper page margins and spacing
"""

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None, header_color="1F4E79"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
        set_cell_shading(hdr_cells[i], header_color)

    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Cm(width)

    return table


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    p._p.get_or_add_pPr().append(shading)


def add_body_text(doc, text, bold=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def generate_english_doc(output_path):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading("GangDan-Refined", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Modular Agent-Pipeline Architecture for LLM Knowledge Management")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x3F, 0x51, 0xB5)
    run.italic = True

    add_body_text(doc, "")
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run('GangDan (\u7eb2\u62c5) \u2014 Principled and Accountable')
    run.font.size = Pt(12)
    run.bold = True

    # Section 1
    doc.add_heading("1. Architecture Overview", level=1)
    add_body_text(doc, "GangDan-Refined uses a modular Agent-Pipeline architecture where each function is an independent, composable agent communicating via a standardized JSON protocol (v2.0). The system has four layers:")

    add_styled_table(doc,
        ["Layer", "Components", "Description"],
        [
            ["CLI Tools", "13 gd-* commands", "Thin shells that parse args, call agents, format output"],
            ["Agent Pipeline", "14 agents + Pipeline engine", "Each agent: AgentInput \u2192 AgentOutput. Pipeline composes agents with | operator"],
            ["Core Modules", "config, i18n, errors, llm, storage, search, learning, document, research", "Shared business logic that agents delegate to"],
            ["Web UI", "Flask + 8 blueprints, 156+ routes", "Settings, chat, KB, docs, learning, research, preprint, export"],
        ],
        col_widths=[3, 5, 8])

    # Section 2
    doc.add_heading("2. Agent Protocol v2.0", level=1)
    add_body_text(doc, "Every agent follows the AgentInput \u2192 AgentOutput protocol. This enables Unix pipe composition and Python Pipeline composition.")

    doc.add_heading("2.1 Input Format", level=2)
    add_code_block(doc, """{
  "query": "quantum computing",
  "text": "Optional text content",
  "file_path": "/path/to/file.pdf",
  "data": {"key": "value"},
  "options": {"model": "qwen2.5:7b", "language": "zh"},
  "metadata": {
    "agent": "gd-search",
    "version": "2.0.0",
    "timestamp": "2026-05-22T10:30:00+00:00",
    "pipeline_id": "pipe_12345678"
  }
}""")

    doc.add_heading("2.2 Output Format", level=2)
    add_code_block(doc, """{
  "success": true,
  "data": {"results": [...], "count": 10, "source": "web"},
  "error": null,
  "metadata": {"agent": "gd-search", "version": "2.0.0", ...},
  "protocol_version": "2.0"
}""")

    doc.add_heading("2.3 Pipeline Composition", level=2)
    add_body_text(doc, "Agents can be chained using the Pipeline class or Unix pipes:", bold=True)
    add_code_block(doc, """# Python: Pipeline with | operator
pipeline = Pipeline(SearchAgent()) | SummarizeAgent() | TranslateAgent()
result = pipeline.run(AgentInput(query="neural networks"))

# CLI: Unix pipe chain
gd-search "topic" --json | gd-summarize --stdin --json | gd-translate --stdin --to zh --json""")

    doc.add_heading("2.4 Pipeline Result Tracking", level=2)
    add_body_text(doc, "Each pipeline execution tracks timing, success, and intermediate data for every step:")

    add_styled_table(doc,
        ["Step", "Duration", "Status", "Description"],
        [
            ["gd-search", "450ms", "Success", "Web search completed"],
            ["gd-summarize", "1200ms", "Success", "Text summarized"],
            ["gd-translate", "800ms", "Success", "Translated to zh"],
            ["Total", "2450ms", "Success", "Pipeline complete"],
        ],
        col_widths=[3, 3, 3, 7])

    # Section 3
    doc.add_heading("3. Agent Reference", level=1)
    add_body_text(doc, "All 14 agents follow the BaseAgent abstract class and communicate via AgentInput/AgentOutput:")

    add_styled_table(doc,
        ["Agent", "CLI Command", "Input", "Output (data keys)", "Description"],
        [
            ["ConfigAgent", "gd-config", "options: {action, key, value}", "{config, key, value}", "View/modify configuration"],
            ["ModelsAgent", "gd-models", "options: {provider}", "{models, count}", "List/manage LLM models"],
            ["ChatAgent", "gd-chat", "query, options: {model, system_prompt}", "{response, model}", "Interactive LLM chat"],
            ["SearchAgent", "gd-search", "query, options: {source, max_results}", "{results, count, source}", "Web & academic search"],
            ["SummarizeAgent", "gd-summarize", "text, options: {style}", "{summary, style, original_length}", "Summarize text"],
            ["TranslateAgent", "gd-translate", "text, options: {target_language}", "{translation, source, target}", "Translate text"],
            ["EmbedAgent", "gd-embed", "text, options: {model}", "{embedding, dimension, model}", "Generate embeddings"],
            ["AskAgent", "gd-ask", "query, options: {kb_names}", "{answer, context_used, sources}", "RAG Q&A over KB"],
            ["KBAgent", "gd-kb", "options: {action, name}", "{kbs, name, count}", "Knowledge base CRUD"],
            ["DocsAgent", "gd-docs", "options: {action, source}", "{sources, document}", "Download/index docs"],
            ["ConvertAgent", "gd-convert", "file_path, options: {format}", "{markdown, pages}", "Convert PDF/HTML/TeX"],
            ["ResearchAgent", "gd-research", "query, options: {phases}", "{report, phases}", "Multi-phase research"],
            ["LearnAgent", "gd-learn", "query, options: {feature}", "{questions, lesson, exam}", "Question gen, learning"],
            ["PreprintAgent", "gd-preprint", "query, options: {action}", "{papers, categories}", "Search/convert preprints"],
        ],
        col_widths=[2.5, 2.5, 3.5, 3, 4])

    # Section 4
    doc.add_heading("4. Configuration Groups", level=1)
    add_body_text(doc, "Configuration is organized into 9 dataclass groups instead of a flat 50+ field monolith:")

    add_styled_table(doc,
        ["Group", "Key Fields", "Example"],
        [
            ["proxy", "mode, http, https", 'proxy.mode = "none"'],
            ["llm", "chat_model, embedding_model, ollama_url, ...", 'llm.chat_model = "qwen2.5:7b"'],
            ["storage", "top_k, chunk_size, chunk_overlap, ...", "storage.top_k = 5"],
            ["search", "web_search_engine, research_sources, ...", 'search.web_search_engine = "duckduckgo"'],
            ["learning", "default_question_type, num_questions, ...", 'learning.default_question_type = "mcq"'],
            ["adaptive", "auto_chunk_size, auto_context_length, ...", "adaptive.auto_chunk_size = True"],
            ["document", "pdf_converter, download_retries, ...", 'document.pdf_converter = "marker"'],
            ["ui", "language, theme, ...", 'ui.language = "en"'],
            ["logging", "level, file, ...", 'logging.level = "INFO"'],
        ],
        col_widths=[3, 6, 6])

    # Section 5
    doc.add_heading("5. Module Structure", level=1)

    add_styled_table(doc,
        ["Module", "Files", "Responsibility"],
        [
            ["agents/", "14 agents + protocol + pipeline + registry", "Agent definitions, JSON protocol, Pipeline composition"],
            ["commands/", "12 CLI command modules + common.py", "Argument parsing, output formatting, thin shells"],
            ["core/", "config, i18n, constants, errors, port_utils", "Shared configuration, translations, error hierarchy"],
            ["llm/", "ollama, openai_compat, factory, models", "LLM abstraction, multi-provider support"],
            ["storage/", "chroma_manager, kb_manager, vector_db, conversation, doc_manager", "Persistence: ChromaDB, KB, chat history, docs"],
            ["search/", "web_searcher, research_searcher, adaptive_search, query_expander", "Search backends: web, academic, query refinement"],
            ["learning/", "question_gen, guided, exam, lecture, research, prompts", "Teaching: questions, lessons, exams, lectures"],
            ["document/", "pdf_converter, pdf_downloader, pdf_renamer, preprint/", "Document processing: PDF, preprints, batch"],
            ["research/", "pipeline, models, export", "Deep research: multi-phase pipeline"],
            ["web/", "app.py + 9 route modules (156+ routes)", "Flask web UI with 8 blueprints"],
        ],
        col_widths=[3, 6, 7])

    # Section 6
    doc.add_heading("6. Test Coverage", level=1)

    add_styled_table(doc,
        ["Test File", "Tests", "Focus"],
        [
            ["test_agent_protocol.py", "60", "AgentInput/Output/Metadata, encode/decode, validation"],
            ["test_agent_base_pipeline.py", "19", "BaseAgent ABC, CLI args, output formatting"],
            ["test_agent_whitebox.py", "35+", "All 14 agents with mocked dependencies"],
            ["test_pipeline_e2e.py", "21", "Pipeline composition with real and test agents"],
            ["test_cli_commands.py", "16", "All 12 CLI commands, JSON/text output"],
            ["test_edge_cases.py", "15", "Error handling, whitespace, None values"],
            ["test_config.py", "13", "Config groups, validation, save/load"],
            ["test_llm.py", "12", "Ollama client, factory, errors"],
            ["test_web_routes.py", "44", "All page routes and API endpoints"],
            ["test_architecture.py", "15", "Module structure, imports, conventions"],
            ["Total", "352", "Comprehensive coverage"],
        ],
        col_widths=[5, 1.5, 9])

    # Section 7
    doc.add_heading("7. Quick Start", level=1)
    doc.add_heading("7.1 Installation", level=2)
    add_code_block(doc, """pip install -e .                  # From source
pip install -e ".[search]"          # With web search
pip install -e ".[pdf]"             # With PDF conversion
pip install -e ".[all]"              # All optional dependencies""")

    doc.add_heading("7.2 CLI Usage", level=2)
    add_code_block(doc, """gd-config --json show              # View configuration
gd-models --json                   # List models
gd-search "topic" --json            # Web search
gd-summarize "text" --style bullet --json
gd-translate "Hello" --to zh --json
gd-kb --action list --json         # Knowledge bases
gd-ask "What is RAG?" --kb-names my-kb --json
gd-web --port 8080                 # Web interface""")

    doc.add_heading("7.3 Python API", level=2)
    add_code_block(doc, """from gangdan_refined.agents import SearchAgent, SummarizeAgent
from gangdan_refined.agents.protocol import AgentInput
from gangdan_refined.agents.pipeline import Pipeline

# Single agent
result = SearchAgent().run(AgentInput(query="transformer"))
print(result.data["results"])

# Pipeline composition
pipeline = Pipeline(SearchAgent()) | SummarizeAgent()
result = pipeline.run(AgentInput(query="quantum computing"))
print(result.data["summary"])""")

    # Section 8
    doc.add_heading("8. Internationalization", level=1)
    add_body_text(doc, "437 translation keys across 10 languages, stored externally in core/locales/translations.json:")

    add_styled_table(doc,
        ["Language", "Code", "Coverage"],
        [
            ["Chinese (\u7b80\u4f53)", "zh", "437/437 (100%)"],
            ["English", "en", "437/437 (100%)"],
            ["Japanese (\u65e5\u672c\u8a9e)", "ja", "437/437 (100%)"],
            ["French (Fran\u00e7ais)", "fr", "437/437 (100%)"],
            ["Russian (\u0420\u0443\u0441\u0441\u043a\u0438\u0439)", "ru", "437/437 (100%)"],
            ["German (Deutsch)", "de", "437/437 (100%)"],
            ["Italian (Italiano)", "it", "437/437 (100%)"],
            ["Spanish (Espa\u00f1ol)", "es", "437/437 (100%)"],
            ["Portuguese (Portugu\u00eas)", "pt", "437/437 (100%)"],
            ["Korean (\ud55c\uad6d\uc5b4)", "ko", "437/437 (100%)"],
        ],
        col_widths=[5, 2, 4])

    # Sections 9-10
    doc.add_heading("9. Requirements", level=1)
    for item in [
        "Python 3.10+",
        "Ollama running locally (default http://localhost:11434)",
        "Chat model (e.g. ollama pull qwen2.5)",
        "Embedding model (e.g. ollama pull nomic-embed-text)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("10. License", level=1)
    add_body_text(doc, "GPL-3.0-or-later. See LICENSE file for details.")

    doc.save(output_path)
    print(f"English Word document saved to {output_path}")


def generate_chinese_doc(output_path):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    for i in range(1, 5):
        hs = doc.styles[f"Heading {i}"]
        hs.font.name = "SimHei"
        hs.element.rPr.rFonts.set(qn("w:eastAsia"), "\u9ed1\u4f53")

    title = doc.add_heading("GangDan-Refined \u2014 \u7eb2\u62c5\u91cd\u6784\u7248", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("\u6a21\u5757\u5316 Agent \u7ba1\u7ebf\u67b6\u6784\u7684\u5927\u8bed\u8a00\u6a21\u578b\u77e5\u8bc6\u7ba1\u7406\u5de5\u5177")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x3F, 0x51, 0xB5)
    run.italic = True

    add_body_text(doc, "")
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("\u7eb2\u62c5 \u2014 \u6709\u7eb2\u9886\u6709\u62c5\u5f53")
    run.font.size = Pt(12)
    run.bold = True

    # Section 1
    doc.add_heading("1. \u67b6\u6784\u603b\u89c8", level=1)
    add_body_text(doc, "GangDan-Refined \u91c7\u7528\u6a21\u5757\u5316 Agent \u7ba1\u7ebf\u67b6\u6784\uff0c\u6bcf\u4e2a\u529f\u80fd\u662f\u4e00\u4e2a\u72ec\u7acb\u7684\u53ef\u7ec4\u5408 Agent\uff0c\u901a\u8fc7\u6807\u51c6\u5316 JSON \u534f\u8bae (v2.0) \u901a\u4fe1\u3002\u7cfb\u7edf\u5206\u4e3a\u56db\u5c42\uff1a")

    add_styled_table(doc,
        ["\u5c42\u7ea7", "\u7ec4\u6210", "\u8bf4\u660e"],
        [
            ["CLI \u5de5\u5177", "13 \u4e2a gd-* \u547d\u4ee4", "\u89e3\u6790\u53c2\u6570\u3001\u8c03\u7528 Agent\u3001\u683c\u5f0f\u5316\u8f93\u51fa"],
            ["Agent \u7ba1\u7ebf", "14 \u4e2a Agent + Pipeline \u5f15\u64ce", "AgentInput \u2192 AgentOutput\uff0cPipeline \u7528 | \u8fd0\u7b97\u7b26\u7ec4\u5408"],
            ["\u6838\u5fc3\u6a21\u5757", "config, i18n, errors, llm, storage, search, learning, document, research", "\u5171\u4eab\u4e1a\u52a1\u903b\u8f91\uff0cAgent \u59d4\u6258\u8c03\u7528"],
            ["Web \u754c\u9762", "Flask + 8 \u4e2a\u84dd\u56fe\uff0c156+ \u8def\u7531", "\u8bbe\u7f6e\u3001\u5bf9\u8bdd\u3001\u77e5\u8bc6\u5e93\u3001\u6587\u6863\u3001\u5b66\u4e60\u3001\u7814\u7a76\u3001\u9884\u5370\u672c\u3001\u5bfc\u51fa"],
        ],
        col_widths=[3, 5, 8], header_color="1F4E79")

    # Section 2
    doc.add_heading("2. Agent \u534f\u8bae v2.0", level=1)
    add_body_text(doc, "\u6bcf\u4e2a Agent \u9075\u5faa AgentInput \u2192 AgentOutput \u534f\u8bae\uff0c\u652f\u6301 Unix \u7ba1\u9053\u7ec4\u5408\u548c Python Pipeline \u7ec4\u5408\u3002")

    doc.add_heading("2.1 \u8f93\u5165\u683c\u5f0f", level=2)
    add_code_block(doc, """{
  "query": "\u91cf\u5b50\u8ba1\u7b97",
  "text": "\u53ef\u9009\u7684\u5f85\u5904\u7406\u6587\u672c\u5185\u5bb9",
  "file_path": "/path/to/document.pdf",
  "data": {"key": "value"},
  "options": {"model": "qwen2.5:7b", "language": "zh"},
  "metadata": {
    "agent": "gd-search",
    "version": "2.0.0",
    "timestamp": "2026-05-22T10:30:00+00:00",
    "pipeline_id": "pipe_12345678"
  }
}""")

    doc.add_heading("2.2 \u8f93\u51fa\u683c\u5f0f", level=2)
    add_code_block(doc, """{
  "success": true,
  "data": {"results": [...], "count": 10, "source": "web"},
  "error": null,
  "metadata": {"agent": "gd-search", "version": "2.0.0", ...},
  "protocol_version": "2.0"
}""")

    doc.add_heading("2.3 \u7ba1\u7ebf\u7ec4\u5408", level=2)
    add_body_text(doc, "Agent \u53ef\u4ee5\u901a\u8fc7 Pipeline \u7c7b\u6216 Unix \u7ba1\u9053\u94fe\u63a5\uff1a", bold=True)
    add_code_block(doc, """# Python\uff1a\u4f7f\u7528 | \u8fd0\u7b97\u7b26\u7684\u7ba1\u7ebf
pipeline = Pipeline(SearchAgent()) | SummarizeAgent() | TranslateAgent()
result = pipeline.run(AgentInput(query="\u795e\u7ecf\u7f51\u7edc"))

# CLI\uff1aUnix \u7ba1\u9053\u94fe
gd-search "\u4e3b\u9898" --json | gd-summarize --stdin --json | gd-translate --stdin --to zh --json""")

    doc.add_heading("2.4 \u7ba1\u7ebf\u7ed3\u679c\u8ddf\u8e2a", level=2)
    add_body_text(doc, "\u6bcf\u6b21\u7ba1\u7ebf\u6267\u884c\u8ddf\u8e2a\u6bcf\u4e2a\u6b65\u9aa4\u7684\u65f6\u95f4\u3001\u6210\u529f\u72b6\u6001\u548c\u4e2d\u95f4\u6570\u636e\uff1a")

    add_styled_table(doc,
        ["\u6b65\u9aa4", "\u8017\u65f6", "\u72b6\u6001", "\u8bf4\u660e"],
        [
            ["gd-search", "450ms", "\u6210\u529f", "\u7f51\u7edc\u641c\u7d22\u5b8c\u6210"],
            ["gd-summarize", "1200ms", "\u6210\u529f", "\u6587\u672c\u5df2\u6458\u8981"],
            ["gd-translate", "800ms", "\u6210\u529f", "\u5df2\u7ffb\u8bd1\u4e3a\u4e2d\u6587"],
            ["\u603b\u8ba1", "2450ms", "\u6210\u529f", "\u7ba1\u7ebf\u6267\u884c\u5b8c\u6210"],
        ],
        col_widths=[3.5, 2.5, 2.5, 7], header_color="1F4E79")

    # Section 3
    doc.add_heading("3. Agent \u53c2\u8003", level=1)
    add_body_text(doc, "\u5168\u90e8 14 \u4e2a Agent \u9075\u5faa BaseAgent \u62bd\u8c61\u7c7b\uff0c\u901a\u8fc7 AgentInput/AgentOutput \u901a\u4fe1\uff1a")

    add_styled_table(doc,
        ["Agent", "CLI \u547d\u4ee4", "\u8f93\u5165", "\u8f93\u51fa (data \u952e)", "\u8bf4\u660e"],
        [
            ["ConfigAgent", "gd-config", "options: {action, key, value}", "{config, key, value}", "\u67e5\u770b/\u4fee\u6539\u914d\u7f6e"],
            ["ModelsAgent", "gd-models", "options: {provider}", "{models, count}", "\u5217\u51fa/\u7ba1\u7406 LLM \u6a21\u578b"],
            ["ChatAgent", "gd-chat", "query, options: {model, system_prompt}", "{response, model}", "\u4ea4\u4e92\u5f0f LLM \u5bf9\u8bdd"],
            ["SearchAgent", "gd-search", "query, options: {source, max_results}", "{results, count, source}", "\u7f51\u7edc\u548c\u5b66\u672f\u641c\u7d22"],
            ["SummarizeAgent", "gd-summarize", "text, options: {style}", "{summary, style, original_length}", "\u6587\u672c\u6458\u8981"],
            ["TranslateAgent", "gd-translate", "text, options: {target_language}", "{translation, source, target}", "\u6587\u672c\u7ffb\u8bd1"],
            ["EmbedAgent", "gd-embed", "text, options: {model}", "{embedding, dimension, model}", "\u751f\u6210\u5d4c\u5165\u5411\u91cf"],
            ["AskAgent", "gd-ask", "query, options: {kb_names}", "{answer, context_used, sources}", "\u57fa\u4e8e\u77e5\u8bc6\u5e93\u7684 RAG \u95ee\u7b54"],
            ["KBAgent", "gd-kb", "options: {action, name}", "{kbs, name, count}", "\u77e5\u8bc6\u5e93\u589e\u5220\u6539\u67e5"],
            ["DocsAgent", "gd-docs", "options: {action, source}", "{sources, document}", "\u4e0b\u8f7d/\u7d22\u5f15\u6587\u6863"],
            ["ConvertAgent", "gd-convert", "file_path, options: {format}", "{markdown, pages}", "\u8f6c\u6362 PDF/HTML/TeX"],
            ["ResearchAgent", "gd-research", "query, options: {phases}", "{report, phases}", "\u591a\u9636\u6bb5\u6df1\u5ea6\u7814\u7a76"],
            ["LearnAgent", "gd-learn", "query, options: {feature}", "{questions, lesson, exam}", "\u51fa\u9898\u3001\u5f15\u5bfc\u5b66\u4e60\u3001\u8003\u8bd5"],
            ["PreprintAgent", "gd-preprint", "query, options: {action}", "{papers, categories}", "\u641c\u7d22/\u8f6c\u6362\u9884\u5370\u672c"],
        ],
        col_widths=[2.5, 2.5, 3.5, 3, 3.5], header_color="1F4E79")

    # Section 4
    doc.add_heading("4. \u914d\u7f6e\u5206\u7ec4", level=1)
    add_body_text(doc, "\u914d\u7f6e\u6309 9 \u4e2a dataclass \u5206\u7ec4\uff0c\u66ff\u4ee3\u539f\u5148 50+ \u5b57\u6bb5\u7684\u6241\u5e73\u7ed3\u6784\uff1a")

    add_styled_table(doc,
        ["\u5206\u7ec4", "\u5173\u952e\u5b57\u6bb5", "\u793a\u4f8b"],
        [
            ["proxy", "mode, http, https", 'proxy.mode = "none"'],
            ["llm", "chat_model, embedding_model, ollama_url, ...", 'llm.chat_model = "qwen2.5:7b"'],
            ["storage", "top_k, chunk_size, chunk_overlap, ...", "storage.top_k = 5"],
            ["search", "web_search_engine, research_sources, ...", 'search.web_search_engine = "duckduckgo"'],
            ["learning", "default_question_type, num_questions, ...", 'learning.default_question_type = "mcq"'],
            ["adaptive", "auto_chunk_size, auto_context_length, ...", "adaptive.auto_chunk_size = True"],
            ["document", "pdf_converter, download_retries, ...", 'document.pdf_converter = "marker"'],
            ["ui", "language, theme, ...", 'ui.language = "zh"'],
            ["logging", "level, file, ...", 'logging.level = "INFO"'],
        ],
        col_widths=[3, 6, 6], header_color="1F4E79")

    # Section 5
    doc.add_heading("5. \u6a21\u5757\u7ed3\u6784", level=1)

    add_styled_table(doc,
        ["\u6a21\u5757", "\u6587\u4ef6", "\u804c\u8d23"],
        [
            ["agents/", "14 \u4e2a Agent + \u534f\u8bae + \u7ba1\u7ebf + \u6ce8\u518c\u8868", "Agent \u5b9a\u4e49\u3001JSON \u534f\u8bae\u3001Pipeline \u7ec4\u5408"],
            ["commands/", "12 \u4e2a CLI \u547d\u4ee4\u6a21\u5757 + common.py", "\u53c2\u6570\u89e3\u6790\u3001\u8f93\u51fa\u683c\u5f0f\u5316\u3001\u8584\u58f3\u5c01\u88c5"],
            ["core/", "config, i18n, constants, errors, port_utils", "\u5171\u4eab\u914d\u7f6e\u3001\u7ffb\u8bd1\u3001\u9519\u8bef\u5c42\u6b21"],
            ["llm/", "ollama, openai_compat, factory, models", "LLM \u62bd\u8c61\u5c42\u3001\u591a\u63d0\u4f9b\u5546\u652f\u6301"],
            ["storage/", "chroma_manager, kb_manager, vector_db, conversation, doc_manager", "\u6301\u4e45\u5316\uff1aChromaDB\u3001\u77e5\u8bc6\u5e93\u3001\u5bf9\u8bdd\u5386\u53f2\u3001\u6587\u6863"],
            ["search/", "web_searcher, research_searcher, adaptive_search, query_expander", "\u641c\u7d22\u540e\u7aef\uff1a\u7f51\u7edc\u3001\u5b66\u672f\u3001\u67e5\u8be2\u7cbe\u70bc"],
            ["learning/", "question_gen, guided, exam, lecture, research, prompts", "\u6559\u5b66\uff1a\u51fa\u9898\u3001\u8bfe\u7a0b\u3001\u8003\u8bd5\u3001\u8bb2\u5ea7"],
            ["document/", "pdf_converter, pdf_downloader, pdf_renamer, preprint/", "\u6587\u6863\u5904\u7406\uff1aPDF\u3001\u9884\u5370\u672c\u3001\u6279\u91cf"],
            ["research/", "pipeline, models, export", "\u6df1\u5ea6\u7814\u7a76\uff1a\u591a\u9636\u6bb5\u7ba1\u7ebf"],
            ["web/", "app.py + 9 \u4e2a\u8def\u7531\u6a21\u5757 (156+ \u8def\u7531)", "Flask Web \u754c\u9762\uff0c8 \u4e2a\u84dd\u56fe"],
        ],
        col_widths=[3, 5.5, 7], header_color="1F4E79")

    # Section 6
    doc.add_heading("6. \u6d4b\u8bd5\u8986\u76d6", level=1)

    add_styled_table(doc,
        ["\u6d4b\u8bd5\u6587\u4ef6", "\u6d4b\u8bd5\u6570", "\u91cd\u70b9"],
        [
            ["test_agent_protocol.py", "60", "AgentInput/Output/Metadata\u3001\u7f16\u89e3\u7801\u3001\u9a8c\u8bc1"],
            ["test_agent_base_pipeline.py", "19", "BaseAgent \u62bd\u8c61\u57fa\u7c7b\u3001CLI \u53c2\u6570\u3001\u8f93\u51fa\u683c\u5f0f"],
            ["test_agent_whitebox.py", "35+", "\u5168\u90e8 14 \u4e2a Agent\uff08\u6a21\u62df\u4f9d\u8d56\uff09"],
            ["test_pipeline_e2e.py", "21", "\u7ba1\u7ebf\u7ec4\u5408\uff08\u771f\u5b9e\u548c\u6d4b\u8bd5 Agent\uff09"],
            ["test_cli_commands.py", "16", "\u5168\u90e8 12 \u4e2a CLI \u547d\u4ee4"],
            ["test_edge_cases.py", "15", "\u9519\u8bef\u5904\u7406\u3001\u7a7a\u767d\u8f93\u5165\u3001None \u503c"],
            ["test_config.py", "13", "\u914d\u7f6e\u5206\u7ec4\u3001\u9a8c\u8bc1\u3001\u4fdd\u5b58/\u52a0\u8f7d"],
            ["test_llm.py", "12", "Ollama \u5ba2\u6237\u7aef\u3001\u5de5\u5382\u3001\u9519\u8bef"],
            ["test_web_routes.py", "44", "\u9875\u9762\u8def\u7531\u548c API \u7aef\u70b9"],
            ["test_architecture.py", "15", "\u6a21\u5757\u7ed3\u6784\u3001\u5bfc\u5165\u3001\u89c4\u8303"],
            ["\u603b\u8ba1", "352", "\u5168\u9762\u8986\u76d6"],
        ],
        col_widths=[5, 1.5, 9], header_color="1F4E79")

    # Section 7
    doc.add_heading("7. \u5feb\u901f\u5f00\u59cb", level=1)
    doc.add_heading("7.1 \u5b89\u88c5", level=2)
    add_code_block(doc, """pip install -e .                  # \u4ece\u6e90\u7801
pip install -e ".[search]"          # \u542b\u7f51\u7edc\u641c\u7d22
pip install -e ".[pdf]"             # \u542b PDF \u8f6c\u6362
pip install -e ".[all]"              # \u6240\u6709\u53ef\u9009\u4f9d\u8d56""")

    doc.add_heading("7.2 \u547d\u4ee4\u884c\u4f7f\u7528", level=2)
    add_code_block(doc, """gd-config --json show              # \u67e5\u770b\u914d\u7f6e
gd-models --json                   # \u5217\u51fa\u6a21\u578b
gd-search "\u4e3b\u9898" --json            # \u7f51\u7edc\u641c\u7d22
gd-summarize "\u6587\u672c" --style bullet --json
gd-translate "Hello" --to zh --json
gd-kb --action list --json         # \u77e5\u8bc6\u5e93
gd-ask "\u4ec0\u4e48\u662f RAG\uff1f" --kb-names my-kb --json
gd-web --port 8080                 # Web \u754c\u9762""")

    doc.add_heading("7.3 Python API", level=2)
    add_code_block(doc, """from gangdan_refined.agents import SearchAgent, SummarizeAgent
from gangdan_refined.agents.protocol import AgentInput
from gangdan_refined.agents.pipeline import Pipeline

# \u5355\u4e2a Agent
result = SearchAgent().run(AgentInput(query="transformer"))
print(result.data["results"])

# \u7ba1\u7ebf\u7ec4\u5408
pipeline = Pipeline(SearchAgent()) | SummarizeAgent()
result = pipeline.run(AgentInput(query="\u91cf\u5b50\u8ba1\u7b97"))
print(result.data["summary"])""")

    # Section 8
    doc.add_heading("8. \u56fd\u9645\u5316", level=1)
    add_body_text(doc, "437 \u4e2a\u7ffb\u8bd1\u952e\uff0c\u652f\u6301 10 \u79cd\u8bed\u8a00\uff0c\u5b58\u50a8\u5728\u5916\u90e8 JSON \u6587\u4ef6\u4e2d\uff1a")

    add_styled_table(doc,
        ["\u8bed\u8a00", "\u4ee3\u7801", "\u8986\u76d6\u7387"],
        [
            ["\u7b80\u4f53\u4e2d\u6587", "zh", "437/437 (100%)"],
            ["English", "en", "437/437 (100%)"],
            ["\u65e5\u672c\u8a9e", "ja", "437/437 (100%)"],
            ["Fran\u00e7ais", "fr", "437/437 (100%)"],
            ["\u0420\u0443\u0441\u0441\u043a\u0438\u0439", "ru", "437/437 (100%)"],
            ["Deutsch", "de", "437/437 (100%)"],
            ["Italiano", "it", "437/437 (100%)"],
            ["Espa\u00f1ol", "es", "437/437 (100%)"],
            ["Portugu\u00eas", "pt", "437/437 (100%)"],
            ["\ud55c\uad6d\uc5b4", "ko", "437/437 (100%)"],
        ],
        col_widths=[5, 2, 4], header_color="1F4E79")

    # Sections 9-10
    doc.add_heading("9. \u73af\u5883\u8981\u6c42", level=1)
    for item in [
        "Python 3.10+",
        "Ollama \u672c\u5730\u8fd0\u884c\uff08\u9ed8\u8ba4 http://localhost:11434\uff09",
        "\u804a\u5929\u6a21\u578b\uff08\u5982 ollama pull qwen2.5\uff09",
        "\u5d4c\u5165\u6a21\u578b\uff08\u5982 ollama pull nomic-embed-text\uff09",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("10. \u8bb8\u53ef\u8bc1", level=1)
    add_body_text(doc, "GPL-3.0-or-later\u3002\u8be6\u89c1 LICENSE \u6587\u4ef6\u3002")

    doc.save(output_path)
    print(f"\u4e2d\u6587 Word \u6587\u6863\u5df2\u4fdd\u5b58\u5230 {output_path}")


if __name__ == "__main__":
    output_dir = os.path.dirname(os.path.abspath(__file__))

    en_path = os.path.join(output_dir, "GangDan-Refined-Documentation.docx")
    generate_english_doc(en_path)

    zh_path = os.path.join(output_dir, "GangDan-Refined-\u6587\u6863.docx")
    generate_chinese_doc(zh_path)

    print(f"\nGenerated:\n  EN: {en_path}\n  ZH: {zh_path}")